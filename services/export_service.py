import os
import io
import json
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch

class ExportService:
    def __init__(self):
        # Setup font for PDF
        self.font_path = "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"
        if os.path.exists(self.font_path):
            pdfmetrics.registerFont(TTFont("ZenHei", self.font_path))
            self.pdf_font = "ZenHei"
        else:
            self.pdf_font = "Helvetica"

    def _format_srt_time(self, seconds):
        td = timedelta(seconds=seconds)
        total_seconds = int(td.total_seconds())
        hours = total_seconds // 3600
        minutes = (total_seconds % 3600) // 60
        secs = total_seconds % 60
        millis = int(td.microseconds / 1000)
        return f"{hours:02}:{minutes:02}:{secs:02},{millis:03}"

    def export_srt(self, transcript: str, speakers: list = None) -> str:
        """
        将转写结果转换为 SRT 字幕格式。
        """
        srt_lines = []
        
        if speakers:
            # Flatten all segments from all speakers and sort them by start time
            all_segments = []
            for sp in speakers:
                label = sp.get("matched_name") or sp.get("label", "Unknown")
                for seg in sp.get("segments", []):
                    all_segments.append({
                        "start": seg.get("start", 0),
                        "end": seg.get("end", 0),
                        "text": seg.get("text", ""),
                        "speaker": label
                    })
            all_segments.sort(key=lambda x: x["start"])
            
            for i, seg in enumerate(all_segments):
                srt_lines.append(str(i + 1))
                start_str = self._format_srt_time(seg["start"])
                end_str = self._format_srt_time(seg["end"])
                srt_lines.append(f"{start_str} --> {end_str}")
                srt_lines.append(f"【{seg['speaker']}】{seg['text']}")
                srt_lines.append("")
        else:
            # Fallback: Split transcript by double newlines and assign estimated times
            paragraphs = [p.strip() for p in transcript.split("\n\n") if p.strip()]
            for i, p in enumerate(paragraphs):
                srt_lines.append(str(i + 1))
                start_str = self._format_srt_time(i * 5)
                end_str = self._format_srt_time((i + 1) * 5)
                srt_lines.append(f"{start_str} --> {end_str}")
                srt_lines.append(p)
                srt_lines.append("")
                
        return "\n".join(srt_lines)

    def export_word(self, transcript: str, metadata: dict = None, speakers: list = None) -> bytes:
        """
        导出为 Word 文档 (.docx)。
        """
        doc = Document()
        
        # Metadata / Title Page
        if metadata:
            filename = metadata.get("filename", "Untitled")
            doc.add_heading(filename, 0)
            
            doc.add_paragraph(f"日期: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            if "file_size_mb" in metadata:
                doc.add_paragraph(f"文件大小: {metadata['file_size_mb']} MB")
        else:
            doc.add_heading("转写结果", 0)

        doc.add_page_break()

        # Body
        if speakers:
            all_segments = []
            for sp in speakers:
                label = sp.get("matched_name") or sp.get("label", "Unknown")
                for seg in sp.get("segments", []):
                    all_segments.append({
                        "start": seg.get("start", 0),
                        "end": seg.get("end", 0),
                        "text": seg.get("text", ""),
                        "speaker": label
                    })
            all_segments.sort(key=lambda x: x["start"])
            
            for seg in all_segments:
                p = doc.add_paragraph()
                run = p.add_run(f"【{seg['speaker']}】")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102) # Dark blue
                
                time_str = f" [{self._format_srt_time(seg['start']).replace(',', '.')} --> {self._format_srt_time(seg['end']).replace(',', '.')}] "
                p.add_run(time_str).font.size = Pt(9)
                
                p.add_run("\n" + seg["text"])
        else:
            doc.add_paragraph(transcript)

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    def export_pdf(self, transcript: str, metadata: dict = None, speakers: list = None) -> bytes:
        """
        导出为 PDF。
        """
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        
        # Custom styles for Chinese support
        chinese_style = ParagraphStyle(
            name='Chinese',
            fontName=self.pdf_font,
            fontSize=10,
            leading=14,
            wordWrap='CJK'
        )
        title_style = ParagraphStyle(
            name='ChineseTitle',
            fontName=self.pdf_font,
            fontSize=18,
            leading=22,
            alignment=1, # Center
            spaceAfter=20
        )
        speaker_style = ParagraphStyle(
            name='ChineseSpeaker',
            fontName=self.pdf_font,
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#003366"),
            wordWrap='CJK'
        )

        elements = []
        
        # Title and Metadata
        if metadata:
            elements.append(Paragraph(metadata.get("filename", "Untitled"), title_style))
            elements.append(Paragraph(f"日期: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", chinese_style))
            if "file_size_mb" in metadata:
                elements.append(Paragraph(f"文件大小: {metadata['file_size_mb']} MB", chinese_style))
        else:
            elements.append(Paragraph("转写结果", title_style))
            
        elements.append(Spacer(1, 0.5 * inch))

        # Body
        if speakers:
            all_segments = []
            for sp in speakers:
                label = sp.get("matched_name") or sp.get("label", "Unknown")
                for seg in sp.get("segments", []):
                    all_segments.append({
                        "start": seg.get("start", 0),
                        "end": seg.get("end", 0),
                        "text": seg.get("text", ""),
                        "speaker": label
                    })
            all_segments.sort(key=lambda x: x["start"])
            
            for seg in all_segments:
                time_str = f"({self._format_srt_time(seg['start']).replace(',', '.')} - {self._format_srt_time(seg['end']).replace(',', '.')})"
                elements.append(Paragraph(f"<b>【{seg['speaker']}】</b> {time_str}", speaker_style))
                elements.append(Paragraph(seg["text"], chinese_style))
                elements.append(Spacer(1, 0.1 * inch))
        else:
            # Handle text with newlines
            for line in transcript.split("\n"):
                if line.strip():
                    elements.append(Paragraph(line, chinese_style))
                else:
                    elements.append(Spacer(1, 0.1 * inch))

        doc.build(elements)
        return buffer.getvalue()
